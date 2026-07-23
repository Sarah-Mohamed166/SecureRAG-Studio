import re

from pathlib import Path


from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models.document import Document


class DocumentLoader:
    """
    Loads supported document types and returns a list of Document objects.
    Each PDF page becomes a separate Document to preserve page numbers.
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt",
        ".md",
    }

    def _clean_text(self, text: str) -> str:
        """
        Normalize whitespace extracted from documents.
        """
        text = re.sub(r"\s+", " ", text)
        return text.strip()

    def load(self, file_path: str) -> list[Document]:
        """
        Detect file type and dispatch to the appropriate loader.
        """

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{extension}'. "
                f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        if extension == ".pdf":
            return self._load_pdf(path)

        if extension == ".docx":
            return self._load_docx(path)

        return self._load_text(path)

    def _load_pdf(self, path: Path) -> list[Document]:
        """
        Extract text page by page from a PDF.
        """

        reader = PdfReader(path)

        documents = []

        for page_number, page in enumerate(reader.pages, start=1):

            text = page.extract_text()

            if text:
                text = self._clean_text(text)

            if not text:
                continue

            documents.append(
                Document(
                    text=text,
                    filename=path.name,
                    page=page_number,
                )
            )

        return documents

    def _load_docx(self, path: Path) -> list[Document]:
        """
        Extract all text from a DOCX file.
        """

        doc = DocxDocument(path)

        paragraphs = [
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        ]

        text = "\n".join(paragraphs)

        return [
            Document(
                text=text,
                filename=path.name,
            )
        ]

    def _load_text(self, path: Path) -> list[Document]:
        """
        Load TXT or Markdown files.
        """

        with open(path, "r", encoding="utf-8") as file:
            text = file.read()

        return [
            Document(
                text=text,
                filename=path.name,
            )
        ]